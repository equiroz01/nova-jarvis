"""Document worker — generate markdown or Word documents."""

import asyncio
import logging

from .base import llm_generate, get_workspace_dir, save_to_brain, parse_json_response

logger = logging.getLogger(__name__)

OUTLINE_SYSTEM = """You are a document architect. Given a document request, create a detailed outline.
Return ONLY a JSON array of objects with this structure:
[{"title": "Section Title", "description": "What to cover in 1-2 sentences"}]
Create 4-8 sections. No explanations, no markdown, just the JSON array.
Write section titles and descriptions in the SAME LANGUAGE as the request."""

SECTION_SYSTEM = """You are a professional document writer. Write detailed content for one section of a document.
Requirements:
- Write in the SAME LANGUAGE as the request
- Use professional, clear tone
- Use markdown formatting (bold, bullets, sub-headers with ###)
- Write 150-400 words per section
- Return ONLY the section content, no extra commentary"""


async def document_worker(task, update_progress) -> str:
    """Generate a markdown or Word document.

    Pipeline: detect format → outline → generate sections → assemble → save
    """
    task_id = task.id
    desc_lower = task.description.lower()

    # Step 1: Detect output format
    want_docx = any(kw in desc_lower for kw in ["word", "docx", ".docx", "documento word"])
    fmt = "DOCX" if want_docx else "Markdown"

    await update_progress(
        task_id, "task_update",
        progress=5, progress_text=f"Formato: {fmt}",
    )

    # Step 2: Generate outline
    await update_progress(task_id, "task_update", progress=10, progress_text="Generando estructura...")

    sections = await _generate_outline(task.description)
    n = len(sections)
    logger.info(f"Document {task_id[:8]}: outline with {n} sections")

    await update_progress(
        task_id, "task_update",
        progress=15, progress_text=f"Estructura: {n} secciones",
    )

    # Step 3: Generate content per section
    section_contents = []
    for i, section in enumerate(sections):
        progress = 20 + int(55 * i / n)
        section_title = section.get("title", f"Section {i+1}")
        await update_progress(
            task_id, "task_update",
            progress=progress, progress_text=f"Escribiendo: {section_title}",
        )

        content = await _generate_section(task.title, task.description, section)
        section_contents.append({"title": section_title, "content": content})

        await asyncio.sleep(0.2)

    # Step 4: Assemble document
    await update_progress(task_id, "task_update", progress=80, progress_text="Ensamblando documento...")

    workspace = get_workspace_dir(task_id)
    markdown = _assemble_markdown(task.title, section_contents)

    if want_docx:
        docx_path = workspace / "document.docx"
        _create_docx(task.title, section_contents, docx_path)
        logger.info(f"Document {task_id[:8]}: DOCX saved to {docx_path}")

    md_path = workspace / "document.md"
    md_path.write_text(markdown, encoding="utf-8")

    # Step 5: Save summary to brain
    await update_progress(task_id, "task_update", progress=90, progress_text="Guardando...")

    summary = f"Document: {task.title} ({n} sections, {fmt})"
    save_to_brain(f"Doc: {task.title}", summary, category="facts")

    result = markdown
    if want_docx:
        result += f"\n\n---\nDocumento Word guardado en: `{docx_path}`"

    return result


async def _generate_outline(description: str) -> list[dict]:
    """Generate document outline via LLM."""
    try:
        response = await llm_generate(description, system=OUTLINE_SYSTEM)
        sections = parse_json_response(response)
        if isinstance(sections, list) and len(sections) >= 2:
            return sections[:8]
    except Exception as e:
        logger.warning(f"Outline generation failed: {e}")

    # Fallback: simple 3-section structure (Spanish, matching user's primary language)
    return [
        {"title": "Introduccion", "description": "Vision general del tema"},
        {"title": "Desarrollo", "description": "Contenido principal y analisis"},
        {"title": "Conclusion", "description": "Resumen y proximos pasos"},
    ]


async def _generate_section(doc_title: str, doc_description: str, section: dict) -> str:
    """Generate content for a single section."""
    prompt = (
        f"Document: {doc_title}\n"
        f"Topic: {doc_description}\n"
        f"Section: {section.get('title', '')}\n"
        f"Guidance: {section.get('description', '')}"
    )
    try:
        return await llm_generate(prompt, system=SECTION_SYSTEM)
    except Exception as e:
        logger.warning(f"Section generation failed: {e}")
        return f"*Content generation failed for this section: {e}*"


def _assemble_markdown(title: str, sections: list[dict]) -> str:
    """Combine sections into a full markdown document."""
    parts = [f"# {title}\n"]
    for s in sections:
        parts.append(f"\n## {s['title']}\n\n{s['content']}\n")
    return "\n".join(parts)


def _create_docx(title: str, sections: list[dict], output_path):
    """Create a Word document using python-docx."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError:
        logger.error("python-docx not installed, cannot generate DOCX")
        raise RuntimeError("python-docx no está instalado. Instalar con: pip install python-docx")

    doc = Document()
    doc.add_heading(title, level=1)

    for s in sections:
        doc.add_heading(s["title"], level=2)
        for para in s["content"].split("\n"):
            text = para.strip()
            if not text:
                continue
            # Strip markdown markers for docx
            text = text.replace("**", "").replace("*", "")
            # Sub-headers
            if text.startswith("### "):
                doc.add_heading(text[4:], level=3)
            elif text.startswith("## "):
                doc.add_heading(text[3:], level=3)
            # Bullet points
            elif text.startswith("- ") or text.startswith("* "):
                doc.add_paragraph(text[2:], style="List Bullet")
            else:
                doc.add_paragraph(text)

    doc.save(str(output_path))
